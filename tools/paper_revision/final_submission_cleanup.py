from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


MAIN_SRC = Path(r"C:/Users/haoha/Desktop/chu_ieee_zhgeng_polished_final_ieee_zh_polished_logic_refined_condensed_submission_refined_final_optimized_related_work_strengthened_balanced_9p_cited_polished_wacv_reframed_slim_main_reviewer_aligned_clean.docx")
SUPP_SRC = Path(r"C:/Users/haoha/Desktop/chu_ieee_zhgeng_polished_final_ieee_zh_polished_logic_refined_condensed_submission_refined_final_optimized_related_work_strengthened_balanced_9p_cited_polished_wacv_reframed_supplementary_reviewer_aligned.docx")

MAIN_DST = MAIN_SRC.with_name(MAIN_SRC.stem + "_finalcheck.docx")
SUPP_DST = SUPP_SRC.with_name(SUPP_SRC.stem + "_finalcheck.docx")


def nonempty_paragraphs(doc):
    return [p for p in doc.paragraphs if p.text.strip()]


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def insert_heading_before(paragraph, text, style_name="Heading 2"):
    new_p = paragraph.insert_paragraph_before(text, style=style_name)
    return new_p


def replace_exact(paragraphs, old, new):
    for p in paragraphs:
        if p.text.strip() == old:
            p.text = new
            return True
    return False


def merge_reference_block(paragraphs, start_prefix, stop_prefix):
    start_idx = None
    stop_idx = None
    for i, p in enumerate(paragraphs):
        txt = p.text.strip()
        if txt.startswith(start_prefix):
            start_idx = i
        elif start_idx is not None and txt.startswith(stop_prefix):
            stop_idx = i
            break
    if start_idx is None or stop_idx is None:
        return
    merged = " ".join(p.text.strip() for p in paragraphs[start_idx:stop_idx])
    paragraphs[start_idx].text = merged
    for p in paragraphs[start_idx + 1:stop_idx]:
        remove_paragraph(p)


def main():
    main_doc = Document(str(MAIN_SRC))
    supp_doc = Document(str(SUPP_SRC))

    main_paras = nonempty_paragraphs(main_doc)

    # Add missing section headings if absent.
    discussion_para = next(p for p in main_paras if p.text.strip().startswith("综合跨数据集比较、消融实验"))
    insert_heading_before(discussion_para, "讨论")
    conclusion_para = next(p for p in nonempty_paragraphs(main_doc) if p.text.strip().startswith("本文提出了一种关键点置信度引导的鲁棒骨架手语识别方法"))
    insert_heading_before(conclusion_para, "结论")

    main_paras = nonempty_paragraphs(main_doc)

    # Remove duplicate conclusion paragraph if present.
    seen_conclusion = False
    for p in list(main_paras):
        txt = p.text.strip()
        if txt.startswith("本文提出了一种关键点置信度引导的鲁棒骨架手语识别方法"):
            if seen_conclusion:
                remove_paragraph(p)
            seen_conclusion = True

    main_paras = nonempty_paragraphs(main_doc)

    # Renumber main tables sequentially.
    replacements = {
        "表 3 和表 4 分别给出了": "表 2 和表 3 分别给出了",
        "表 3 SLR500 上的标准测试结果比较": "表 2 SLR500 上的标准测试结果比较",
        "表 4 NMFs-CSL 上的标准测试结果比较": "表 3 NMFs-CSL 上的标准测试结果比较",
        "表 5 给出了不同模块配置下的结果。": "表 4 给出了不同模块配置下的结果。",
        "表 5 WLASL100 上的模块消融结果（含质量加权融合与固定平均融合对照）": "表 4 WLASL100 上的模块消融结果（含质量加权融合与固定平均融合对照）",
        "表 5 中不同模块配置": "表 4 中不同模块配置",
        "综合表 5 与图 3": "综合表 4 与图 3",
        "表 6 进一步用于区分": "表 5 进一步用于区分",
        "表 6 一致性训练的比较结果（WLASL100，关节点单流）": "表 5 一致性训练的比较结果（WLASL100，关节点单流）",
        "表 6 表明": "表 5 表明",
        "结果如表 8 所示。": "结果如表 6 所示。",
        "表 8 WLASL100 上的关键点缺失鲁棒性结果": "表 6 WLASL100 上的关键点缺失鲁棒性结果",
        "结果如表 9 所示。": "结果如表 7 所示。",
        "表 9 WLASL100 上的坐标噪声鲁棒性结果": "表 7 WLASL100 上的坐标噪声鲁棒性结果",
        "结合表 11 的比较结果": "结合表 8 的比较结果",
        "表 11 与强骨架基线的统一扰动对比结果（关节点单流 Top-1）": "表 8 与强骨架基线的统一扰动对比结果（关节点单流 Top-1）",
        "表 11 在统一 DSTA-SLR [7] 主干": "表 8 在统一 DSTA-SLR [7] 主干",
    }
    for p in main_paras:
        txt = p.text
        for old, new in replacements.items():
            if old in txt:
                txt = txt.replace(old, new)
        p.text = txt

    # Merge split references.
    main_paras = nonempty_paragraphs(main_doc)
    merge_reference_block(main_paras, "[15]", "[16]")
    main_paras = nonempty_paragraphs(main_doc)
    merge_reference_block(main_paras, "[16]", "[17]")

    main_doc.save(str(MAIN_DST))

    # Supplementary cleanup and renumbering.
    supp_paras = nonempty_paragraphs(supp_doc)

    # Remove duplicate overhead heading/intro.
    overhead_heading_seen = False
    overhead_intro_seen = False
    for p in list(nonempty_paragraphs(supp_doc)):
        txt = p.text.strip()
        if txt == "开销分析":
            if overhead_heading_seen:
                remove_paragraph(p)
            overhead_heading_seen = True
        if txt == "为保持主文聚焦，参数量与 FLOPs 的完整表格移至补充材料，主文仅保留其关键结论。":
            if overhead_intro_seen:
                remove_paragraph(p)
            overhead_intro_seen = True

    supp_paras = nonempty_paragraphs(supp_doc)
    supp_replacements = {
        "表 2 进一步汇报了": "补充表 S1 进一步汇报了",
        "表 2 MSASL 各子集上的标准测试结果比较": "补充表 S1 MSASL 各子集上的标准测试结果比较",
        "表 7 将原始置信度": "补充表 S2 将原始置信度",
        "表 7 简单置信度基线的相对比较结果（20 个 epoch 短程验证）": "补充表 S2 简单置信度基线的相对比较结果（20 个 epoch 短程验证）",
        "表 7 的结果表明": "补充表 S2 的结果表明",
        "结果见表 10。": "结果见补充表 S3。",
        "表 10 不同样本质量区间上的准确率": "补充表 S3 不同样本质量区间上的准确率",
        "表 13 表明": "补充表 S4 表明",
        "表 13 与公开 DSTA-SLR 强基线的统一扰动对比结果（WLASL100，关节点单流 Top-1）": "补充表 S4 与公开 DSTA-SLR 强基线的统一扰动对比结果（WLASL100，关节点单流 Top-1）",
        "表 12 参数量与计算开销分析（WLASL100，关节点单流主干）": "补充表 S5 参数量与计算开销分析（WLASL100，关节点单流主干）",
    }
    for p in supp_paras:
        txt = p.text
        for old, new in supp_replacements.items():
            if old in txt:
                txt = txt.replace(old, new)
        p.text = txt

    supp_doc.save(str(SUPP_DST))
    print(f"Created main: {MAIN_DST}")
    print(f"Created supp: {SUPP_DST}")


if __name__ == "__main__":
    main()
