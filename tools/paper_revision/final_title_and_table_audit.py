from pathlib import Path
import csv

from docx import Document


DOCX_PATH = Path(r"C:/Users/haoha/Desktop/ieee_chu_submission_final_wacv_cn_submission.docx")
AUDIT_PATH = Path(r"C:/Users/haoha/Documents/New project/docs/paper_revision/table_text_consistency_check.md")

NEW_TITLE = "关键点置信度引导的鲁棒骨架孤立词手语识别"
NEW_KEYWORDS = "关键词：孤立词手语识别；骨架表示；关键点置信度；鲁棒性建模；退化观测"


def update_title_and_keywords(doc: Document):
    title_para = doc.paragraphs[0]
    keywords_para = doc.paragraphs[5]

    if title_para.runs:
        title_para.runs[0].text = NEW_TITLE
        for run in title_para.runs[1:]:
            run.text = ""
    else:
        title_para.text = NEW_TITLE

    if keywords_para.runs:
        keywords_para.runs[0].text = NEW_KEYWORDS
        for run in keywords_para.runs[1:]:
            run.text = ""
    else:
        keywords_para.text = NEW_KEYWORDS


def build_audit():
    lines = [
        "# 表格与正文一致性核对",
        "",
        f"- 文档：`{DOCX_PATH}`",
        "",
        "## 核对结论",
        "",
        "- 表 1：正文中“WLASL100 为 78.86/79.59，高于 BEST，低于 DSTA-SLR 与 NLA-SLR”与表内数值一致。",
        "- 表 2-3：正文中“高于 ST-GCN、低于 DSTA-SLR 及更强方法”与表内数值一致。",
        "- 表 4：正文中的单流与四流描述已与当前表内数值对齐，单流列为 63.42 / 65.10 / 65.10 / 67.28 / 65.60。",
        "- 表 5：正文中的 63.42→65.60、8.39→48.49、67.95 / 59.40 / 58.56 / 63.32 与表内数值一致。",
        "- 表 6：正文中的“30% 关键点缺失下为 47.99%”与表内数值一致。",
        "- 表 7：正文中的“20 像素噪声下为 45.13%”与表内数值一致。",
        "- 表 8：正文未逐项复述全部数值，但其定性描述“结构增强主攻缺失，一致性增强补足强噪声”与表内趋势一致。",
        "",
        "## 仍需注意的点",
        "",
        "- 当前不存在明显的正文-表格数字冲突。",
        "- 但 `Table 4` 与 `Table 5/8` 使用的是不同展示口径：表 4 侧重单流/四流消融，表 5/8 侧重统一受控扰动协议下的鲁棒性结果。数字已不冲突，但投稿时仍建议在答辩中准备解释这一口径差异。",
        "- `Table 5` 中基线四流结果显著低于单流结果这一现象仍会吸引审稿人注意，虽然不构成当前稿件的数字错误，但建议在答辩中强调融合协议与实验目的不同。",
    ]
    AUDIT_PATH.write_text("\n".join(lines), encoding="utf-8")


def main():
    doc = Document(str(DOCX_PATH))
    update_title_and_keywords(doc)
    doc.save(str(DOCX_PATH))
    build_audit()
    print(DOCX_PATH)
    print(AUDIT_PATH)


if __name__ == "__main__":
    main()
