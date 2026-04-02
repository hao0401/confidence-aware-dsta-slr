from copy import deepcopy
from pathlib import Path
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


MAIN_SRC = Path(r"C:/Users/haoha/Desktop/chu_ieee_zhgeng_polished_final_ieee_zh_polished_logic_refined_condensed_submission_refined_final_optimized_related_work_strengthened_balanced_9p_cited_polished_wacv_reframed_slim_main.docx")
SUPP_SRC = Path(r"C:/Users/haoha/Desktop/chu_ieee_zhgeng_polished_final_ieee_zh_polished_logic_refined_condensed_submission_refined_final_optimized_related_work_strengthened_balanced_9p_cited_polished_wacv_reframed_supplementary.docx")

MAIN_DST = MAIN_SRC.with_name(MAIN_SRC.stem + "_reviewer_aligned.docx")
SUPP_DST = SUPP_SRC.with_name(SUPP_SRC.stem + "_reviewer_aligned.docx")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
ns = {
    "w": W_NS,
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
}


def iter_blocks(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield ("p", Paragraph(child, doc))
        elif child.tag.endswith("}tbl"):
            yield ("t", Table(child, doc))


def append_block(doc, block):
    body = doc._element.body
    sect = body.sectPr
    new_el = deepcopy(block._element)
    if sect is None:
        body.append(new_el)
    else:
        body.insert(body.index(sect), new_el)


def insert_before_paragraph(target_para, block):
    parent = target_para._element.getparent()
    parent.insert(parent.index(target_para._element), deepcopy(block._element))


def remove_block(block):
    parent = block._element.getparent()
    if parent is not None:
        parent.remove(block._element)


def nonempty_paragraphs(doc):
    return [p for p in doc.paragraphs if p.text.strip()]


def xml_counts(path: Path):
    with zipfile.ZipFile(path) as z:
        root = ET.fromstring(z.read("word/document.xml"))
    return {
        "tables": len(root.findall(".//w:tbl", ns)),
        "math": len(root.findall(".//m:oMath", ns)) + len(root.findall(".//m:oMathPara", ns)),
        "drawings": len(root.findall(".//w:drawing", ns)),
    }


def main():
    main_doc = Document(str(MAIN_SRC))
    supp_doc = Document(str(SUPP_SRC))

    main_paras = nonempty_paragraphs(main_doc)
    # Tighten claims and explain why controlled robustness focuses on WLASL100.
    main_paras[59 - 1].text = (
        "据此，实验从以下六个方面对所提方法进行系统评估：1) 跨数据集结果比较；2) 模块消融；3) 一致性训练分析；"
        "4) 简单置信度基线比较；5) 与强骨架基线的统一扰动对比；6) 参数量与计算开销分析。需要说明的是，"
        "表 7 所示的简单置信度基线采用 20 个 epoch 的短程验证协议，以控制额外实验成本并突出不同置信度设置的相对排序；"
        "其余主实验均采用完整训练设置。整体而言，这些实验围绕三个问题展开：标准测试下是否保持竞争性性能、退化输入下是否获得显著鲁棒性增益、"
        "以及这些增益是否来自可靠性建模本身而非更大的模型容量。受主文篇幅限制，受控鲁棒性分析聚焦于更具挑战性的 WLASL100，"
        "而跨数据集标准结果用于验证该框架在英文和中文词级手语基准上的适用性。"
    )
    main_paras[95 - 1].text = (
        "综合跨数据集比较、消融实验、简单置信度基线、统一扰动对比和开销分析可以看出，本文方法的主要价值在于将关键点可靠性系统性地引入骨架识别流程。"
        "标准测试结果表明，该方法在 WLASL100、WLASL2000、SLR500 和 NMFs-CSL 等设置下保持了较强的竞争力，但并未在所有基准上稳定超过 DSTA-SLR [7] 等强基线。"
        "固定平均融合控制实验表明，主文中的性能提升并非仅由质量加权融合规则带来，而主要源于输入质量建模及其与结构模块的协同作用；"
        "统一扰动对比则进一步显示，本文结构增强版本在关键点缺失场景下具有明显优势，但在强坐标噪声场景下仍存在差距。"
        "总体而言，本文工作最核心的贡献不在于提出更大的骨架主干，而在于证明：将关键点置信度系统性地贯穿到识别链路中，能够以极低开销显著提升退化骨架输入下的鲁棒性。"
        "与公开 checkpoint 的直接对比和更多补充结果见 supplementary。"
    )
    main_paras[99 - 1].text = (
        "本文提出了一种关键点置信度引导的鲁棒骨架手语识别方法，将关键点置信度作为观测质量先验，统一引入输入编码、空间关系建模、时间特征修复、多流融合和训练阶段的一致性约束。"
        "标准测试结果表明，本文方法在 WLASL100 上取得 78.86%/79.59% 的 P-I/P-C，在 SLR500 上取得 95.58% 的 Top-1，并在多个数据集上保持了具有竞争力的标准测试性能。"
        "更关键的是，固定平均融合控制实验表明，当前多流性能提升在不同融合口径下保持一致，说明主要收益来自质量建模和结构增强本身，而非质量加权融合规则的偶然偏置；"
        "与公开 DSTA-SLR [7] 强基线的统一扰动对比进一步显示，本文结构增强版本在关键点缺失场景下具有显著优势，而在强坐标噪声场景下仍存在改进空间。"
        "总体而言，本文工作证明了一个更具一般性的结论：相较于单纯增强骨架主干容量，系统利用关键点可靠性信息是一条更低开销、且更适合退化观测场景的鲁棒识别路径。"
    )

    # Pull back Chinese benchmark evidence from supplementary.
    target_para = next(p for p in main_doc.paragraphs if p.text.strip() == "模块消融分析")
    supp_blocks = list(iter_blocks(supp_doc))
    collect = False
    blocks_to_insert = []
    for kind, block in supp_blocks:
        txt = block.text.strip() if kind == "p" else None
        if kind == "p" and txt == "SLR500 和 NMFs-CSL":
            collect = True
        if collect:
            if kind == "p" and txt == "置信度基线比较":
                break
            blocks_to_insert.append(block)
    for block in blocks_to_insert:
        insert_before_paragraph(target_para, block)

    # Move overhead table and caption to supplementary to offset added space.
    main_blocks = list(iter_blocks(main_doc))
    move_to_supp = []
    capture = False
    for kind, block in main_blocks:
        txt = block.text.strip() if kind == "p" else None
        if kind == "p" and txt == "表 12 参数量与计算开销分析（WLASL100，关节点单流主干）":
            capture = True
        if capture:
            move_to_supp.append(block)
            if kind == "t":
                break
    if move_to_supp:
        append_block(supp_doc, supp_doc.add_paragraph("开销分析"))
        append_block(supp_doc, supp_doc.add_paragraph("为保持主文聚焦，参数量与 FLOPs 的完整表格移至补充材料，主文仅保留其关键结论。"))
        for block in move_to_supp:
            append_block(supp_doc, block)
            remove_block(block)

    # Rewrite overhead paragraph in main to avoid missing table reference.
    for p in nonempty_paragraphs(main_doc):
        if p.text.strip().startswith("表 12 表明，所提结构模块几乎不增加模型规模。"):
            p.text = (
                "开销分析表明，所提结构模块几乎不增加模型规模。与基线相比，完整模型的参数量仅由 7.219534 M 增加到 7.219868 M，"
                "FLOPs 仅由 6.088505 G 增加到 6.088977 G，增量分别为 0.000334 M 和 0.000472 G。"
                "这说明，本文方法的鲁棒性收益并非来自显著更大的模型容量，而是来自对观测质量的显式建模；同时，可靠性感知一致性仅作用于训练阶段，因此推理阶段的额外成本主要来自结构模块本身，部署负担较小。"
            )
            break

    main_doc.save(str(MAIN_DST))
    supp_doc.save(str(SUPP_DST))

    if xml_counts(MAIN_DST)["drawings"] != xml_counts(MAIN_SRC)["drawings"]:
        raise RuntimeError("Unexpected drawing count mismatch in main")
    print(f"Created main: {MAIN_DST}")
    print(f"Created supplementary: {SUPP_DST}")


if __name__ == "__main__":
    main()
