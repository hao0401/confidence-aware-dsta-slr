from __future__ import annotations

from pathlib import Path

from docx import Document


SRC = Path(r"C:\Users\haoha\Desktop\ieee_chu_submission_final.docx")


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def main() -> None:
    doc = Document(str(SRC))

    # Merge split reference [15]
    merged_15 = (
        '[15] L. T. Woods and Z. A. Rana, "Modelling sign language '
        'with encoder-only transformers and human pose estimation keypoint data," '
        'Mathematics, vol. 11, no. 9, p. 2129, 2023.'
    )
    doc.paragraphs[157].text = merged_15
    for idx in [160, 159, 158]:
        remove_paragraph(doc.paragraphs[idx])

    # After removing [15] continuation lines, [16] starts at the original anchor paragraph.
    merged_16 = (
        '[16] M. Pu, M. K. Lim, and C. Y. Chong, "Siformer: '
        'Feature-isolated transformer for efficient skeleton-based sign language recognition," '
        'in Proc. 32nd ACM Int. Conf. Multimedia (MM), 2024, pp. 9387-9396.'
    )
    doc.paragraphs[158].text = merged_16
    for idx in [161, 160, 159]:
        remove_paragraph(doc.paragraphs[idx])

    doc.save(str(SRC))
    print(SRC)


if __name__ == "__main__":
    main()
